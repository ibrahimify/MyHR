from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

AUTHOR = "Muhammad Ibrahim Shoeb"
SUPERVISOR = "Dr. Husam Al-Maghoosi"
DATE = "17 May 2026"


def setup(doc):
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        field(footer, "PAGE")
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 30, "030213"),
        ("Heading 1", 18, "030213"),
        ("Heading 2", 14, "111827"),
        ("Heading 3", 12, "1f2937"),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "CodeBlock" not in doc.styles:
        style = doc.styles.add_style("CodeBlock", 1)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        style.font.size = Pt(8.5)


def field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def update_fields_on_open(doc):
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings.element.append(update)


def title_page(doc, title, subtitle, audience):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MyHR")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("2563eb")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("4b5563")
    doc.add_paragraph()
    table(doc, ["Field", "Value"], [
        ["Audience", audience],
        ["Project", "BME Project Lab desktop HR management system"],
        ["Author", AUTHOR],
        ["Supervisor", SUPERVISOR],
        ["Date", DATE],
    ])
    doc.add_page_break()


def toc(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph("If Word does not populate the table automatically, right-click it and choose Update Field.")
    note.runs[0].italic = True
    doc.add_page_break()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()
    return t


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Screenshot placeholder: {label}]")
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("6b7280")


def code(doc, text):
    for line in text.strip().splitlines():
        doc.add_paragraph(line.rstrip(), style="CodeBlock")
    doc.add_paragraph()


def build_user_guide():
    doc = Document()
    setup(doc)
    update_fields_on_open(doc)
    title_page(doc, "User Guide", "Day-one manual for Admin and HR Officer users", "Admin and HR Officer")
    toc(doc)

    doc.add_heading("1. Purpose and Roles", level=1)
    doc.add_paragraph("MyHR is an offline desktop HR system for employee records, organization hierarchy, promotion races, commendations, sanctions, salary increments, imports, exports, backups, and audit review.")
    table(doc, ["Role", "Main Access"], [
        ["Administrator", "All pages, Settings, salary ranges, promotion rules, annual increments, user management, export, backup, password changes."],
        ["HR Officer", "Dashboard, Employees, Hierarchy, Promotions, Commendations, Sanctions, Audit Log, and Import Data. Settings is admin-only."],
    ])

    doc.add_heading("2. Log In, Log Out, and Switch Language", level=1)
    numbers(doc, [
        "Open MyHR.",
        "Choose English, Hungarian, or Arabic on the login screen.",
        "Enter username and password.",
        "Press Sign In.",
    ])
    table(doc, ["Default Account", "Username", "Password"], [
        ["Administrator", "admin", "admin123"],
        ["HR Officer", "hr_officer", "hr123"],
    ])
    bullets(doc, [
        "Use Logout at the bottom of the left sidebar to sign out and return to the login screen.",
        "Arabic uses right-to-left layout for the application session.",
        "Only application UI text is translated. Client-entered data, names, notes, emails, employee IDs, job titles, CSV headers, audit descriptions, and database values remain exactly as entered.",
    ])
    placeholder(doc, "Login screen with language selector")

    doc.add_heading("3. Dashboard", level=1)
    bullets(doc, [
        "Review total employees, pending promotions, active sanctions, commendations this month, recent activity, and upcoming promotions.",
        "Use Review in Salary Increments Due to approve one or all yearly salary increments.",
        "Approved increments update employee salary, create salary history, and write an audit log entry.",
    ])
    placeholder(doc, "Dashboard and salary increment review dialog")

    doc.add_heading("4. Employee Management", level=1)
    doc.add_heading("Add employees", level=2)
    numbers(doc, [
        "Open Employees and press Add Employee.",
        "Fill personal details, work details, degree, position, join date, base salary, status, org unit, and manager.",
        "Choose BSc, MSc, PhD, or Other. BSc starts at L7, MSc at L6, PhD at L5, and Other uses the Other/Misc track.",
        "Watch the live salary range warning. The record cannot be saved if salary is outside the admin-defined range.",
        "Save the employee. MyHR generates an employee ID and logs the action.",
    ])
    doc.add_heading("Other or miscellaneous employees", level=2)
    bullets(doc, [
        "Use Other for janitors, guards, cleaners, painters, window cleaners, and similar roles.",
        "The displayed level is Other/Misc, not L7 to L1.",
        "They are assigned to the OTHERS branch and can report only to the OTHERS head, CEO, or a configured employee responsible for Other staff.",
        "They receive annual increments, but they do not have main level promotions, commendations, or sanctions.",
    ])
    doc.add_heading("View and edit employees", level=2)
    bullets(doc, [
        "Use search and filters on the employee list.",
        "Open the profile with View. Personal Details shows the main race and sub-race immediately.",
        "Promotion History shows promotions and annual increment milestones as an audit-style history.",
        "Use Edit to update the employee. Use Delete only when a record was created by mistake because related employee history is removed too.",
    ])
    placeholder(doc, "Employee list, add form, and employee profile")

    doc.add_heading("5. Organization Hierarchy", level=1)
    table(doc, ["Type", "Rule"], [
        ["Organization", "Root only. Only one organization is allowed."],
        ["Division", "Must be under an organization."],
        ["Department", "Must be under a division."],
        ["Unit", "Must be under a department."],
        ["Team", "Must be under a unit."],
        ["Position", "Must be under a team."],
        ["OTHERS", "Special branch for Other/Misc employees."],
    ])
    bullets(doc, [
        "Use Add Unit or Add Department to create nodes.",
        "Unavailable unit types are disabled until their required parent exists.",
        "Assign a Head or In-Charge where needed.",
        "Click populated units or teams to view employees in that scope.",
    ])
    placeholder(doc, "Organization hierarchy and OTHERS branch")

    doc.add_heading("6. Promotion Race and Sub-Race", level=1)
    doc.add_paragraph("The main race is the level movement, for example L7 to L6. The sub-race is the yearly checkpoint inside the race, for example L7.1 and L7.2.")
    table(doc, ["Item", "Meaning"], [
        ["Main race", "Configured in Settings as level-to-level months, for example L7 to L6 in 36 months."],
        ["Sub-race", "Yearly milestones inside the main race. Completed steps turn green in the profile UI."],
        ["Annual increment", "Applied yearly at the employee's anniversary using the configured percentage or fixed amount."],
        ["Promotion increase", "Applied on promotion using the next level's configured promotion salary increase."],
        ["Commendation", "Reduces months remaining for standard employees."],
        ["Active sanction", "Adds months remaining until resolved."],
    ])
    numbers(doc, [
        "Open Promotions.",
        "Review eligible employees and soon-eligible employees.",
        "Press Approve for an eligible employee.",
        "Confirm the salary and level change.",
        "MyHR updates the employee, creates promotion history, and writes the audit log.",
    ])
    placeholder(doc, "Promotion page and profile race timeline")

    doc.add_heading("7. Commendations", level=1)
    bullets(doc, [
        "Open Commendations and choose Issue Commendation.",
        "Choose Single Employee or Team Award.",
        "Choose Category 1 (-1 month), Category 2 (-3 months), or Category 3 (-6 months).",
        "Each standard employee can receive up to three commendations in the current role.",
        "Team awards skip employees who already reached the limit.",
    ])
    placeholder(doc, "Commendation issue form and history")

    doc.add_heading("8. Sanctions", level=1)
    bullets(doc, [
        "Open Sanctions and choose Issue Sanction.",
        "Select employee, sanction type, reason, and delay from 1 to 12 months.",
        "Active sanctions delay the promotion race.",
        "Use Active Sanctions to mark a sanction resolved. Resolved sanctions stay in history but stop adding active delay.",
    ])
    placeholder(doc, "Sanction issue form, active sanctions, and history")

    doc.add_heading("9. Import, Export, and Backup", level=1)
    bullets(doc, [
        "Import Data accepts CSV and XLSX files.",
        "Required headers are first_name, last_name, department, degree, position, join_date, and base_salary.",
        "Optional headers include division, unit, team, work_email, work_phone, personal_email, phone, address, status, manager_work_email, commendation_months, active_sanction_months, and sanction_type.",
        "The preview marks valid rows and rows needing review before anything is written.",
        "Admin can export employees and create a SQLite database backup from Settings > Database.",
    ])
    placeholder(doc, "Import preview, export dialog, and backup dialog")

    doc.add_heading("10. User Management and Audit", level=1)
    bullets(doc, [
        "Admin uses Settings > User Management to change password and manage system users.",
        "The admin account can be edited for handover.",
        "HR Officer accounts can be created, edited, deactivated, and reactivated.",
        "Deactivation is a soft delete: the HR account cannot log in, but the record remains for audit history.",
        "Audit Log shows user identity as username: full name, for example admin: John Doe.",
        "Past audit entries keep their original username and full-name snapshot even if the account changes later.",
    ])
    placeholder(doc, "User Management tab and Audit Log")

    doc.save(DOCS / "MyHR_User_Guide.docx")


def build_developer_guide():
    doc = Document()
    setup(doc)
    update_fields_on_open(doc)
    title_page(doc, "Developer Guide", "Technical handover guide for MyHR", "Future developer / thesis continuation")
    toc(doc)

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph("MyHR is a standalone offline desktop HRMS written in Python. It uses PySide6 for the UI, SQLAlchemy for persistence, SQLite as a local database, qtawesome icons, dictionary-based internationalization, and Qt QSettings for simple company branding.")

    doc.add_heading("2. Tech Stack and Why", level=1)
    table(doc, ["Layer", "Technology", "Reason"], [
        ["Language", "Python 3", "Rapid development and strong desktop/database libraries."],
        ["Desktop UI", "PySide6 / Qt", "Native widgets, tables, dialogs, layout management, and RTL support."],
        ["Database", "SQLite", "Offline single-file storage with simple backup."],
        ["ORM", "SQLAlchemy", "Structured schema, relationships, and safer queries."],
        ["Icons", "qtawesome / Font Awesome 5", "Consistent professional UI icons."],
        ["Excel import", "openpyxl with fallback XML parser", "Supports XLSX onboarding files."],
        ["Settings", "Qt QSettings", "Stores company name and subtitle outside schema changes."],
        ["Docs", "python-docx", "Generates required Word handover documents."],
    ])

    doc.add_heading("3. Folder Structure", level=1)
    code(doc, """
MyHR/
  main.py
  requirements.txt
  README.md
  docs/
  scripts/seed_demo_company.py
  src/
    core/app_settings.py
    core/i18n.py
    database/models.py
    database/connection.py
    ui/login_window.py
    ui/main_window.py
    ui/styles.py
    ui/pages/dashboard.py
    ui/pages/employees.py
    ui/pages/hierarchy.py
    ui/pages/promotions.py
    ui/pages/commendations.py
    ui/pages/sanctions.py
    ui/pages/audit_log.py
    ui/pages/import_data.py
    ui/pages/settings.py
""")
    table(doc, ["Area", "Responsibility"], [
        ["main.py", "Application entry point, database initialization, global Qt styling, login startup."],
        ["src/core", "Internationalization and persisted company settings."],
        ["src/database/models.py", "SQLAlchemy table definitions and relationships."],
        ["src/database/connection.py", "Engine/session setup, seed defaults, auth, ID generation, race math, salary increments, audit helper, Other/Misc helpers."],
        ["src/ui/main_window.py", "Sidebar navigation, role checks, page cache/refresh, logout."],
        ["src/ui/pages", "Feature page implementations."],
    ])

    doc.add_heading("4. Database Schema", level=1)
    table(doc, ["Table", "Columns", "Relationships and Notes"], [
        ["system_user", "id, username, password_hash, role, full_name, is_active, created_at, last_login", "Admin and HR accounts. is_active=False is soft delete."],
        ["org_unit", "id, name, unit_type, parent_id, head_employee_id, created_at, updated_at", "Self-referencing hierarchy. Optional head_employee_id to employee."],
        ["title", "id, name, label, degree_requirement, base_salary_min, base_salary_max, currency, annual_increment_type, annual_increment_value, promotion_salary_increase_pct, created_at, updated_at", "Salary level and increment configuration. Includes L7-L1 and Other."],
        ["employee", "id, employee_id, first_name, last_name, date_of_birth, phone, personal_email, address, degree, work_email, work_phone, position, join_date, base_salary, status, annual_increment_last_applied, profile_picture, linkedin_url, org_unit_id, title_id, reports_to_id, created_at, updated_at", "Core employee record. Links to org_unit, title, and manager employee."],
        ["promotion_rule", "id, from_title_id, to_title_id, base_months, is_active, created_at, updated_at", "Configured level transition duration."],
        ["promotion_history", "id, employee_id, from_title_id, to_title_id, approved_by_id, basis, months_taken, notes, promoted_at", "Immutable promotion record."],
        ["commendation", "id, commendation_ref, title, description, category, months_impact, is_team_award, issued_by_id, issued_at", "Award record. Affects standard promotion race."],
        ["commendation_employee", "commendation_id, employee_id", "Join table for single and team commendations."],
        ["sanction", "id, sanction_ref, employee_id, sanction_type, reason, delay_months, issued_by_id, issued_at, resolved_at, is_resolved", "Disciplinary record. Active sanctions delay standard promotion race."],
        ["salary_increment_history", "id, employee_id, approved_by_id, salary_before, salary_after, increment_type, increment_value, applied_at, notes", "Annual increment audit history."],
        ["audit_log", "id, performed_by_id, performed_by_username, performed_by_name, action, target_table, target_id, description, before_value, after_value, performed_at", "Immutable action stream with user identity snapshot."],
    ])

    doc.add_heading("5. Business Logic", level=1)
    doc.add_heading("Authentication and access control", level=2)
    doc.add_paragraph("verify_login hashes the password, checks active system_user rows, updates last_login, and returns a detached UserSession. main_window.py hides and blocks Settings for HR Officers. Employees are data subjects only and never log in.")
    doc.add_heading("Internationalization", level=2)
    doc.add_paragraph("i18n.py stores English, Hungarian, and Arabic dictionaries. set_language changes the session language; t(key, **kwargs) resolves UI text. Arabic switches the Qt layout direction to right-to-left. Internal enum values, audit action codes, import headers, export headers, and client-entered data are not translated.")
    doc.add_heading("Audit immutability", level=2)
    doc.add_paragraph("log_action creates audit_log rows with performed_by_id plus performed_by_username and performed_by_name snapshots. The Audit Log page displays username: full name. Renaming a user does not rewrite old audit rows.")
    doc.add_heading("User management and soft delete", level=2)
    doc.add_paragraph("Admin can edit the admin account and manage HR accounts. HR deletion sets is_active=False instead of deleting the row, preserving historical audit joins and preventing unknown users in the audit UI.")
    doc.add_heading("Other/Misc employees", level=2)
    doc.add_paragraph("degree_to_title_name maps Other to title Other. ensure_others_org_unit creates the OTHERS branch. valid_other_manager_ids limits managers to the OTHERS head, CEO, or dedicated Other employee handlers. Other employees receive annual increments but have no main promotion track and are excluded from commendation and sanction selection.")
    doc.add_heading("Hierarchy rules", level=2)
    doc.add_paragraph("The hierarchy enforces exactly one organization and parent order organization -> division -> department -> unit -> team -> position. Disabled dropdown options and save-time validation prevent invalid structures and parent cycles.")

    doc.add_heading("6. Promotion Race Formula", level=1)
    code(doc, """
race_start = last_promotion_date or employee.join_date
months_elapsed = full_months_between(race_start, today)
commendation_reduction = sum(abs(months_impact)) for commendations since race_start
sanction_addition = sum(delay_months) for unresolved sanctions since race_start
raw_remaining = base_months - months_elapsed - commendation_reduction + sanction_addition
months_remaining = max(0, raw_remaining)
eligible = months_remaining == 0
progress_pct = clamp((months_elapsed + commendation_reduction - sanction_addition) / base_months * 100, 0, 100)
""")
    doc.add_paragraph("calculate_sub_race builds yearly L7.1/L7.2 style milestones plus the final promotion step. Other employees get ongoing Other.1, Other.2, and so on for annual increments only. Promotion approval applies the next title's promotion_salary_increase_pct to current salary and creates PromotionHistory. Annual increments are separate and create SalaryIncrementHistory.")

    doc.add_heading("7. CSV Import Validation Rules", level=1)
    bullets(doc, [
        "Required: first_name, last_name, department, degree, position, join_date, base_salary.",
        "Header aliases normalize common names such as email to work_email and salary to base_salary.",
        "Degree must be BSc, MSc, PhD, or Other.",
        "Join date must be YYYY-MM-DD and not in the future.",
        "Base salary must be numeric and greater than zero, then must pass configured salary range validation during import.",
        "Status defaults to active and must be active, inactive, or on_leave.",
        "Work email is format-checked, unique against the database, and unique inside the file.",
        "manager_work_email is format-checked and wired after rows are created.",
        "commendation_months must be 0, 1, 3, or 6.",
        "active_sanction_months must be 0 to 12.",
        "sanction_type accepts verbal warning, written warning, suspension, or final warning when sanction months exist.",
    ])

    doc.add_heading("8. Functional Requirements", level=1)
    bullets(doc, [
        "Authenticate Admin and HR Officer users.",
        "Support English, Hungarian, and Arabic UI with RTL for Arabic.",
        "Manage employees, hierarchy, promotions, commendations, sanctions, audit logs, and imports.",
        "Configure salary ranges, promotion rules, and annual increment rules.",
        "Support Other/Misc employees with OTHERS hierarchy and annual increments only.",
        "Calculate promotion eligibility live and show main/sub-race status.",
        "Approve promotions and annual increments with history records.",
        "Create, edit, deactivate, and reactivate HR accounts.",
        "Export employees and back up the SQLite database.",
        "Log important actions with immutable audit identity snapshots.",
    ])

    doc.add_heading("9. Non-Functional Requirements", level=1)
    bullets(doc, [
        "Offline-first operation without a runtime server.",
        "Professional and consistent PySide6 UI.",
        "Role-based access control.",
        "Data integrity through ORM relationships and validation before writes.",
        "Auditability with user snapshots.",
        "Localization without changing client data or internal contracts.",
        "Simple backup and export for operational safety.",
    ])

    doc.add_heading("10. Resource Requirements", level=1)
    table(doc, ["Resource", "Requirement"], [
        ["OS", "Windows development environment; PySide6 can be tested on other desktop OSes."],
        ["Python", "Python 3.12+ recommended."],
        ["RAM", "4 GB minimum recommended."],
        ["Storage", "Project files, SQLite database, exports, and backups."],
        ["Dependencies", "PySide6, SQLAlchemy, qtawesome, openpyxl, python-docx."],
        ["Network", "Only needed to install dependencies. Runtime is offline."],
    ])

    doc.add_heading("11. PlantUML", level=1)
    diagrams = {
        "Use Case Diagram": """
@startuml
left to right direction
actor Admin
actor "HR Officer" as HR
rectangle MyHR {
  usecase "Log in" as Login
  usecase "Switch language" as Lang
  usecase "Manage employees" as Emp
  usecase "Build hierarchy" as Hier
  usecase "Approve promotions" as Promo
  usecase "Issue commendations" as Comm
  usecase "Issue sanctions" as Sanc
  usecase "Resolve sanctions" as Resolve
  usecase "Import employees" as Import
  usecase "View audit log" as Audit
  usecase "Approve salary increments" as Inc
  usecase "Configure settings" as Settings
  usecase "Manage user accounts" as Users
  usecase "Export and backup" as Backup
}
Admin --> Login
HR --> Login
Admin --> Lang
HR --> Lang
Admin --> Emp
HR --> Emp
Admin --> Hier
HR --> Hier
Admin --> Promo
HR --> Promo
Admin --> Comm
HR --> Comm
Admin --> Sanc
HR --> Sanc
Admin --> Resolve
HR --> Resolve
Admin --> Import
HR --> Import
Admin --> Audit
HR --> Audit
Admin --> Inc
HR --> Inc
Admin --> Settings
Admin --> Users
Admin --> Backup
@enduml
""",
        "Login Sequence": """
@startuml
actor User
participant LoginWindow
participant "verify_login" as Auth
database SQLite
participant MainWindow
User -> LoginWindow: enter credentials
LoginWindow -> Auth: verify_login(username, password)
Auth -> SQLite: query active system_user
SQLite --> Auth: user or none
Auth -> SQLite: update last_login
Auth --> LoginWindow: UserSession
LoginWindow -> MainWindow: open app
@enduml
""",
        "Employee Creation Sequence": """
@startuml
actor "Admin/HR" as User
participant EmployeesPage
participant EmployeeDialog
participant "connection helpers" as Helpers
database SQLite
User -> EmployeesPage: Add Employee
EmployeesPage -> EmployeeDialog: open form
User -> EmployeeDialog: save employee
EmployeeDialog -> Helpers: degree_to_title_name and validate_salary_for_title
Helpers --> EmployeeDialog: level and validation
EmployeeDialog -> SQLite: insert employee
EmployeeDialog -> SQLite: insert audit_log
SQLite --> EmployeesPage: refresh
@enduml
""",
        "Promotion Approval Sequence": """
@startuml
actor "Admin/HR" as User
participant PromotionsPage
participant "calculate_months_remaining" as Race
database SQLite
User -> PromotionsPage: approve promotion
PromotionsPage -> Race: calculate race
Race -> SQLite: read employee, rule, commendations, sanctions
Race --> PromotionsPage: eligible result
PromotionsPage -> SQLite: update title and salary
PromotionsPage -> SQLite: insert promotion_history
PromotionsPage -> SQLite: insert audit_log
@enduml
""",
        "Commendation Issuance Sequence": """
@startuml
actor "Admin/HR" as User
participant CommendationsPage
participant "cap checks" as Cap
database SQLite
User -> CommendationsPage: issue commendation
CommendationsPage -> Cap: can_receive_commendation
Cap -> SQLite: count current-role commendations
Cap --> CommendationsPage: allowed recipients
CommendationsPage -> SQLite: insert commendation and recipients
CommendationsPage -> SQLite: insert audit_log
@enduml
""",
        "Sanction Issuance Sequence": """
@startuml
actor "Admin/HR" as User
participant SanctionsPage
database SQLite
User -> SanctionsPage: issue sanction
SanctionsPage -> SanctionsPage: validate employee, type, reason, delay
SanctionsPage -> SQLite: insert sanction
SanctionsPage -> SQLite: insert audit_log
@enduml
""",
        "Salary Increment Approval Sequence": """
@startuml
actor "Admin/HR" as User
participant DashboardPage
participant "apply_salary_increment" as Increment
database SQLite
User -> DashboardPage: approve increment
DashboardPage -> Increment: apply_salary_increment(employee_id, user_id)
Increment -> SQLite: read employee and title
Increment -> Increment: calculate salary_after
Increment -> SQLite: update employee salary
Increment -> SQLite: insert salary_increment_history
Increment -> SQLite: insert audit_log
@enduml
""",
    }
    for title, text in diagrams.items():
        doc.add_heading(title, level=2)
        code(doc, text)

    doc.add_heading("12. Future Scope for Thesis Semester", level=1)
    bullets(doc, [
        "Package the desktop app into a user-friendly installer.",
        "Add automated tests for race math, import validation, access control, user management, and audit snapshots.",
        "Add richer reports for yearly promotions, salary budget impact, sanctions, commendations, and departments.",
        "Add scheduled backup reminders and backup health checks.",
        "Introduce Alembic or another migration tool if schema changes become frequent.",
        "Consider encrypted backups or protected local storage for sensitive HR data.",
        "Extend roles only if the organization needs more than Admin and HR Officer.",
    ])

    doc.save(DOCS / "MyHR_Developer_Guide.docx")


if __name__ == "__main__":
    build_user_guide()
    build_developer_guide()
    print(DOCS / "MyHR_User_Guide.docx")
    print(DOCS / "MyHR_Developer_Guide.docx")
